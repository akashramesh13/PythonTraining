from docx import Document

# using SQLAlchemy not FLASK-SQLALCHEMY
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy import Column, Integer, String, Table
from sqlalchemy import create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from sqlalchemy.sql.schema import MetaData

# in memory database to avoid primary key consequences
engine = create_engine('sqlite:///:memory:')
Base = declarative_base(engine)
Session = sessionmaker(bind=engine)
session = Session()
meta = MetaData()

# Object for mapping to table
records_table = Table(
    'records', meta,
    Column("id", Integer, primary_key=True),
    Column("qty", String),
    Column("desc", String)
)
# creating the table
meta.create_all(engine)


document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
# making connection to db
conn = engine.connect()
# sample data
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)
# writing to db
for qty, id, desc in records:
    record = records_table.insert().values(id=id, qty=qty, desc=desc)
    conn.execute(record)


table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

# reading all from database
r = records_table.select()
res = conn.execute(r)

# for each row in res perform the write operation
for qty, id, desc in res:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()
document.save('demo.docx')
